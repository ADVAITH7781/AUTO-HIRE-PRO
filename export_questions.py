import json
import os
from docx import Document

# Config
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
QUESTIONS_DIR = os.path.join(BASE_DIR, "questions")

def export_all_questions():
    if not os.path.exists(QUESTIONS_DIR):
        print("❌ No 'questions' directory found.")
        return

    files = [f for f in os.listdir(QUESTIONS_DIR) if f.endswith(".json")]
    
    if not files:
        print("ℹ️ No question files found to export.")
        return

    print(f"📂 Found {len(files)} question banks. Exporting to Word...")

    for filename in files:
        try:
            job_id = filename.replace(".json", "")
            json_path = os.path.join(QUESTIONS_DIR, filename)
            
            with open(json_path, "r") as f:
                questions = json.load(f)
            
            if not questions or not isinstance(questions, list):
                print(f"⚠️ Skipping {filename}: Invalid or empty JSON.")
                continue

            # Create Doc
            doc = Document()
            doc.add_heading(f'Question Bank: {job_id}', 0)
            
            for i, q in enumerate(questions):
                q_text = q.get('q', 'No Question Text')
                options = q.get('options', [])
                answer = q.get('answer', 'Unknown')
                
                doc.add_paragraph(f"Q{i+1}. {q_text}", style='List Number')
                for opt in options:
                    doc.add_paragraph(str(opt), style='List Bullet')
                
                p = doc.add_paragraph()
                runner = p.add_run(f"Correct Answer: {answer}")
                runner.bold = True
                doc.add_paragraph("-" * 50)
                
            docx_filename = f"{job_id}_Question_Bank.docx"
            docx_path = os.path.join(BASE_DIR, docx_filename)
            doc.save(docx_path)
            print(f"✅ Exported: {docx_filename}")
            
        except Exception as e:
            print(f"❌ Failed to export {filename}: {e}")

if __name__ == "__main__":
    export_all_questions()
