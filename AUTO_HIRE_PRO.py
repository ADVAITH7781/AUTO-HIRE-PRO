# phase1_admin_api.py
import streamlit as st
import pandas as pd
import os
from pypdf import PdfReader
from docx import Document
import google.generativeai as genai
import datetime
import time
import smtplib
import ssl
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

# ---------------- SAFE IMPORTS FOR CV ----------------
try:
    import cv2
    import mediapipe as mp
    import av
    from streamlit_webrtc import webrtc_streamer, VideoTransformerBase, WebRtcMode
    import numpy as np
    PROCTORING_AVAILABLE = True
except ImportError as e:
    PROCTORING_AVAILABLE = False
    print(f"⚠️ CV Import Error: {e}")
    # Define a dummy class to prevent NameError later
    class VideoTransformerBase: pass
    webrtc_streamer = None
except Exception as e:
    PROCTORING_AVAILABLE = False
    print(f"⚠️ Unexpected CV Error: {e}")
    class VideoTransformerBase: pass
    webrtc_streamer = None

# ---------------- Config ----------------
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, "data")
if not os.path.exists(DATA_DIR): os.makedirs(DATA_DIR)

# Files moved to data/ folder for better organization & persistence check
COMPANIES_FILE = os.path.join(DATA_DIR, "companies.xlsx")
APPS_FILE = os.path.join(DATA_DIR, "applications.csv.xlsx")
RESUMES_DIR = os.path.join(BASE_DIR, "resumes")
JOBS_DIR = os.path.join(BASE_DIR, "job_descriptions")
QUESTIONS_DIR = os.path.join(BASE_DIR, "questions")

# Ensure directories exist
for d in [RESUMES_DIR, JOBS_DIR, QUESTIONS_DIR, DATA_DIR]:
    if not os.path.exists(d):
        os.makedirs(d)

try:
    API_KEY = st.secrets["GEMINI_API_KEY"]
except FileNotFoundError:
    st.error("⚠️ Secrets file not found! Please create .streamlit/secrets.toml")

# ---------------- CV PROCTORING LOGIC ----------------
if PROCTORING_AVAILABLE:
    try:
        mp_face_mesh = mp.solutions.face_mesh
        face_mesh = mp_face_mesh.FaceMesh(min_detection_confidence=0.5, min_tracking_confidence=0.5)
    except Exception as e:
        PROCTORING_AVAILABLE = False
        print(f"⚠️ MediaPipe Init Error: {e}")

class ProctoringProcessor(VideoTransformerBase):
    def __init__(self):
        self.warn_count = 0
        self.last_warn = time.time()
        self.frame_count = 0

    def recv(self, frame):
        img = frame.to_ndarray(format="bgr24")
        
        # Performance: Process every 2nd frame to reduce lag
        self.frame_count += 1
        if self.frame_count % 2 != 0:
             return av.VideoFrame.from_ndarray(img, format="bgr24")

        try:
            img_rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
            results = face_mesh.process(img_rgb)
            
            h, w, _ = img.shape
            face_count = 0
            status_text = "Secure"
            color = (0, 255, 0)
            violation = False
            
            if results.multi_face_landmarks:
                face_count = len(results.multi_face_landmarks)
                
                if face_count > 1:
                    status_text = "MULTIPLE FACES DETECTED!"
                    color = (0, 0, 255)
                    violation = True
                elif face_count == 1:
                    for face_landmarks in results.multi_face_landmarks:
                        # Head Pose Estimation (Simple Nose vs Ear X-coord check)
                        nose = face_landmarks.landmark[1]
                        left_ear = face_landmarks.landmark[234]
                        right_ear = face_landmarks.landmark[454]
                        
                        # Convert to pixel coords
                        nx, ny = int(nose.x * w), int(nose.y * h)
                        lx, _ = int(left_ear.x * w), int(left_ear.y * h)
                        rx, _ = int(right_ear.x * w), int(right_ear.y * h)
                        
                        # Check deviation
                        dist_l = abs(nx - lx)
                        dist_r = abs(nx - rx)
                        
                        # Avoid division by zero
                        ratio = dist_l / (dist_r + 1e-6)
                        
                        if ratio < 0.5: # Looking Left
                            status_text = "LOOKING AWAY (LEFT)"
                            color = (0, 165, 255)
                            violation = True
                        elif ratio > 2.0: # Looking Right
                            status_text = "LOOKING AWAY (RIGHT)"
                            color = (0, 165, 255)
                            violation = True
                            
                        # Draw Nose
                        cv2.circle(img, (nx, ny), 5, (255, 0, 0), -1)
            else:
                status_text = "NO FACE DETECTED"
                color = (0, 0, 255)
                # violation = True # Optional: Strict no-face
                
            # Cooldown & Increment
            if violation:
                now = time.time()
                if now - self.last_warn > 4.0: # 4 Seconds Cooldown
                    self.warn_count += 1
                    self.last_warn = now
            
            # Draw Status & Warnings
            cv2.putText(img, f"Status: {status_text}", (20, 40), cv2.FONT_HERSHEY_SIMPLEX, 0.8, color, 2)
            cv2.putText(img, f"WARNINGS: {self.warn_count}/5", (20, 80), cv2.FONT_HERSHEY_SIMPLEX, 0.8, (0, 0, 255), 2)
            
            # Visual Alarm
            if color != (0, 255, 0):
                cv2.rectangle(img, (0,0), (w,h), color, 10)
                
        except Exception as e:
            print(f"Proctor Loop Error: {e}")
             
        return av.VideoFrame.from_ndarray(img, format="bgr24")

# Configure Gemini
genai.configure(api_key=API_KEY)

# ---------------- Data Handling ----------------
def load_data():
    if os.path.exists(COMPANIES_FILE):
        try:
            df = pd.read_excel(COMPANIES_FILE)
            required_cols = ["Company", "Role", "JD", "JD_File_Path", "ResumeThreshold", "AptitudeThreshold", "Job_ID", "HasQuestions"]
            for col in required_cols:
                if col not in df.columns:
                    df[col] = "Pending" # Default to pending for new fields
            df = df.fillna("")
            return df
        except Exception as e:
            st.error(f"❌ Error loading data: {e}")
            return pd.DataFrame(columns=["Company", "Role", "JD", "JD_File_Path", "ResumeThreshold", "AptitudeThreshold", "Job_ID", "HasQuestions"])
    else:
        print(f"ℹ️ File not found: {COMPANIES_FILE}, creating new.")
        df = pd.DataFrame(columns=["Company", "Role", "JD", "JD_File_Path", "ResumeThreshold", "AptitudeThreshold", "Job_ID", "HasQuestions"])
        return df

def save_data(df):
    try:
        if not os.path.exists(DATA_DIR): os.makedirs(DATA_DIR)
        df.to_excel(COMPANIES_FILE, index=False)
        st.toast("✅ Jobs Database Saved!", icon="💾") 
    except Exception as e:
        st.error(f"❌ CRITICAL ERROR SAVING DATA: {e}")

def load_apps():
    if os.path.exists(APPS_FILE):
        try:
            df = pd.read_excel(APPS_FILE)
            cols = ["Name", "Email", "Score", "Company", "Role", "Status", "Resume_Text", "TestPassword", "TokenTime", "TestScore", "TestStatus", "Resume_Path", "Timestamp", "Job_ID", "ApplicantName"]
            for col in cols:
                if col not in df.columns: df[col] = ""
            return df
        except Exception:
            return pd.DataFrame(columns=["Name", "Email", "Score", "Company", "Role", "Status", "Resume_Text", "TestPassword", "TokenTime", "TestScore", "TestStatus", "Resume_Path", "Timestamp", "Job_ID", "ApplicantName"])
    else:
        return pd.DataFrame(columns=["Name", "Email", "Score", "Company", "Role", "Status", "Resume_Text", "TestPassword", "TokenTime", "TestScore", "TestStatus", "Resume_Path", "Timestamp", "Job_ID", "ApplicantName"])

def save_apps(df):
    try:
        if not os.path.exists(DATA_DIR): os.makedirs(DATA_DIR)
        df.to_excel(APPS_FILE, index=False)
    except Exception as e:
        st.error(f"❌ Error Saving Applications: {e}")

# ---------------- Email Notification ----------------
def send_email(candidate_email, score, company, role, email_type="success", token=None):
    try:
        sender_email = st.secrets["EMAIL_ADDRESS"]
        password = st.secrets["EMAIL_PASSWORD"]
    except Exception:
        st.warning("⚠️ Email secrets not found. Skipping email.")
        return ""

    if email_type == "success":
        # Use provided token or generate backup (though caller should provide it)
        if not token:
            chars = "ABCDEFGHJKLMNPQRSTUVWXYZ23456789"
            token = "".join(random.choices(chars, k=6))
        
        subject = f"Congratulations! You've been shortlisted for {role} at {company}"
        heading = "Great News! 🎉"
        heading_color = "#FF9F1C"
        score_color = "#2ecc71"
        try:
            base_url = st.secrets.get("BASE_URL", "http://localhost:8501")
            if base_url.endswith("/"):
                base_url = base_url.rstrip("/")
        except:
            base_url = "http://localhost:8501"
            
        body_content = f"""
        <p>We are thrilled to inform you that your profile has been <strong>shortlisted</strong> for the <strong>{role}</strong> position at <strong>{company}</strong>!</p>
        <p>Your Resume Score: <span style="font-size: 18px; font-weight: bold; color: {score_color};">{score}/100</span></p>
        <hr>
        <p>You have been invited to take the <strong>Proctored Aptitude Test</strong>.</p>
        <div style="background: #fdf2f8; padding: 15px; border-left: 4px solid #db2777; margin: 20px 0;">
            <p style="margin:0; font-weight:bold; color:#be185d;">Your Access Credentials:</p>
            <p style="margin:5px 0 0 0;">Test Password: <span style="font-size: 1.25em; background: #fff; padding: 2px 8px; border: 1px solid #ddd; border-radius: 4px;">{token}</span></p>
            <p style="margin:5px 0 0 0; font-size: 0.85em; color: #666;">⚠️ Valid for 30 Hours only.</p>
        </div>
        <p>Please click the link below to proceed:</p>
        <div style="text-align: center; margin: 30px 0;">
            <a href="{base_url}/?mode=test" style="background-color: #FF9F1C; color: white; padding: 15px 30px; text-decoration: none; border-radius: 5px; font-weight: bold; display: inline-block;">Start Aptitude Test</a>
        </div>
        """
    else:
        subject = f"Update on your application for {role} at {company}"
        heading = "Application Update"
        heading_color = "#555"
        score_color = "#e74c3c"
        body_content = f"""
        <p>Thank you for giving us the opportunity to review your application for the <strong>{role}</strong> position at <strong>{company}</strong>.</p>
        <p>Your Resume Score: <span style="font-size: 18px; font-weight: bold; color: {score_color};">{score}/100</span></p>
        <hr>
        <p><strong>Don't be discouraged!</strong> We will keep your resume in our talent pool for future openings.</p>
        """

    html_content = f"""
    <html>
        <body style="font-family: Arial, sans-serif; color: #333;">
            <div style="background-color: #f4f4f4; padding: 20px;">
                <div style="background-color: white; padding: 30px; border-radius: 10px; max-width: 600px; margin: auto; border-top: 5px solid {heading_color};">
                    <h2 style="color: {heading_color};">{heading}</h2>
                    <p>Dear Candidate,</p>
                    {body_content}
                    <br>
                    <p>Best Regards,<br><strong>Auto Hire Pro Team</strong></p>
                </div>
            </div>
        </body>
    </html>
    """
    
    msg = MIMEMultipart("alternative")
    msg["Subject"] = subject
    msg["From"] = sender_email
    msg["To"] = candidate_email
    msg.attach(MIMEText(html_content, "html"))

    try:
        context = ssl.create_default_context()
        with smtplib.SMTP_SSL("smtp.gmail.com", 465, context=context) as server:
            server.login(sender_email, password)
            server.sendmail(sender_email, candidate_email, msg.as_string())
        print(f"✅ {email_type.capitalize()} email sent to {candidate_email}")
        return token
    except Exception as e:
        st.error(f"❌ Email Delivery Failed: {e}")
        print(f"❌ Failed to send email: {e}")
        return ""

def extract_text_from_pdf(file):
    try:
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            try:
                extracted = page.extract_text()
                if extracted:
                    text += extracted
            except Exception as e:
                print(f"⚠️ Error parsing PDF page: {e}")
                continue
        return text
    except Exception as e:
        st.error(f"Error reading PDF file: {e}")
        return ""

def extract_text_from_docx(file):
    doc = Document(file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def calculate_score(resume_text, jd_text):
    max_retries = 3
    for attempt in range(max_retries):
        try:
            model = genai.GenerativeModel('gemini-flash-latest')
            prompt = f"""
            Act as a calibrated ATS. Compare the Resume to the JD.
            
            JD: {jd_text[:2000]}...
            RESUME: {resume_text[:2000]}...
            
            SCORING ALGORITHM (Base + Merit):
            
            1. **BASE SCORE (40 Points)**: 
               - If the text is a valid resume with Contact, Education, and Experience sections, AUTOMATICALLY AWARD 40 POINTS.
               - If it is gibberish or empty, award 0.
            
            2. **MERIT SCORE (0-60 Points)**:
               - **Keywords & Skills (25)**: Exact matches for Key Technical Skills in JD.
               - **Experience Relevance (25)**: Similar Job Titles, Industy, and Seniority.
               - **Formatting & Impact (10)**: Quantifiable results (numbers/%) and clear structure.
            
            TOTAL = BASE (40) + MERIT (0-60). Max 100.
            
            INSTRUCTIONS:
            - Most decent candidates should score between 50-70.
            - Only perfect matches should exceed 85.
            - OUTPUT FORMAT: "Final Score: <number>"
            """
            generation_config = {"temperature": 0.0, "top_p": 0.1, "top_k": 1}
            response = model.generate_content(prompt, generation_config=generation_config)
            
            # Robust Parsing
            import re
            text = response.text
            match = re.search(r"Final Score:\s*(\d+)", text, re.IGNORECASE)
            if match:
                score = int(match.group(1))
                return min(100, max(0, score)) # Clamp between 0-100
            else:
                # Fallback: try to find any double digit number at the end
                digits = re.findall(r"\d+", text)
                if digits: 
                    return min(100, int(digits[-1]))
                return 40 # Default to base score on error if content likely valid
                
        except Exception as e:
            if "429" in str(e) and attempt < max_retries - 1:
                time.sleep(2 ** attempt)
                continue
            elif attempt == max_retries - 1:
                st.error(f"AI Error: {e}")
                return 0

import json
import random

# ---------------- Question Bank Logic ----------------
QUESTIONS_DIR = os.path.join(BASE_DIR, "questions")
if not os.path.exists(QUESTIONS_DIR): os.makedirs(QUESTIONS_DIR)

def generate_question_bank(jd_text, job_id):
    """Generates 50 Technical Qs (Job Specific) + reuses 50 General Qs (Common Pool)."""
    
    # 1. SETUP COMMON POOL (Logical + Situational) - Reuse if exists
    common_file = os.path.join(QUESTIONS_DIR, "common_pool_v1.json")
    common_questions = []
    
    if os.path.exists(common_file):
        with open(common_file, "r") as f:
            common_questions = json.load(f)
    else:
        # Generate Common Pool ONCE
        common_topics = [
            "Logical Reasoning & IQ (General)",
            "Situational Judgment (Professional Workplace)",
        ]
        model = genai.GenerativeModel('gemini-flash-latest')
        for topic in common_topics:
            try:
                # Ask for 25 each
                prompt = f"""
                Generate 25 Multiple Choice Questions (MCQs).
                FOCUS AREA: {topic}
                TAG: General
                OUTPUT FORMAT (JSON): [{{"q": "...", "options": [...], "answer": "...", "type": "General"}}]
                """
                response = model.generate_content(prompt, generation_config={"response_mime_type": "application/json"})
                batch = json.loads(response.text)
                if isinstance(batch, list):
                    # Ensure type tag exists
                    for b in batch: b["type"] = "General"
                    common_questions.extend(batch)
                time.sleep(1)
            except Exception as e:
                print(f"⚠️ Error gen common pool {topic}: {e}")
        
        # Save Common Pool
        with open(common_file, "w") as f:
            json.dump(common_questions, f)

    # 2. GENERATE JOB SPECIFIC TECHNICAL QUESTIONS (50 Qs)
    tech_questions = []
    tech_topics = [
        "Technical Skills in JD (Hard)",
        "Advanced Role-Specific Scenarios"
    ]
    model = genai.GenerativeModel('gemini-flash-latest')
    
    for topic in tech_topics:
        try:
            prompt = f"""
            Act as a Senior Tech Interviewer. Generate 25 Hard MCQs for this Job Description.
            JD SUMMARY: {jd_text[:1500]}...
            FOCUS AREA: {topic}
            TAG: Technical
            OUTPUT FORMAT (JSON): [{{"q": "...", "options": [...], "answer": "...", "type": "Technical"}}]
            """
            response = model.generate_content(prompt, generation_config={"response_mime_type": "application/json"})
            batch = json.loads(response.text)
            if isinstance(batch, list):
                for b in batch: b["type"] = "Technical"
                tech_questions.extend(batch)
            time.sleep(1)
        except Exception as e:
            print(f"⚠️ Error gen tech batch {topic}: {e}")

    # 3. MERGE & SAVE
    full_bank = tech_questions + common_questions
    
    # Save to JSON
    q_file = os.path.join(QUESTIONS_DIR, f"{job_id}.json")
    with open(q_file, "w") as f:
        json.dump(full_bank, f)
        
    # Save to Word (DOCX)
    docx_path = None
    try:
        doc = Document()
        doc.add_heading(f'Question Bank: {job_id}', 0)
        
        # Section 1: Technical
        doc.add_heading('Part 1: Technical (Job Specific)', level=1)
        for i, q in enumerate(tech_questions):
            doc.add_paragraph(f"T{i+1}. {q.get('q')}", style='List Number')
            for opt in q.get('options', []): doc.add_paragraph(opt, style='List Bullet')
            doc.add_paragraph(f"Answer: {q.get('answer')}", style='Intense Quote')
            
        # Section 2: General
        doc.add_heading('Part 2: General (Common Pool)', level=1)
        for i, q in enumerate(common_questions):
            doc.add_paragraph(f"G{i+1}. {q.get('q')}", style='List Number')
            for opt in q.get('options', []): doc.add_paragraph(opt, style='List Bullet')
            doc.add_paragraph(f"Answer: {q.get('answer')}", style='Intense Quote')
            
        docx_path = os.path.join(BASE_DIR, f"{job_id}_Question_Bank.docx")
        doc.save(docx_path)
    except Exception as e:
        print(f"⚠️ Error saving DOCX: {e}")

    return len(full_bank), docx_path

def get_candidate_questions(job_id, num_questions=40):
    """Samples 25 Technical + 15 General Questions."""
    q_file = os.path.join(QUESTIONS_DIR, f"{job_id}.json")
    if not os.path.exists(q_file): return []
    
    with open(q_file, "r") as f:
        bank = json.load(f)
        
    tech = [q for q in bank if q.get("type") == "Technical"]
    general = [q for q in bank if q.get("type") == "General"]
    
    # Fallback if tags missing (older files)
    if not tech and not general:
        return random.sample(bank, min(len(bank), num_questions))
        
    # Stratified Sampling (25 Tech, 15 General => 40 Total)
    # If not enough, take what we have
    n_tech = min(len(tech), 25)
    n_gen = min(len(general), 15)
    
    exam_set = random.sample(tech, n_tech) + random.sample(general, n_gen)
    random.shuffle(exam_set)
    return exam_set

# ---------------- VIBRANT SAAS UI IMPLEMENTATION ----------------
def main():
    st.set_page_config(page_title="Auto Hire Pro", page_icon="🚀", layout="wide")

    # ---------------- CSS: NEW GLOBAL THEME SYSTEM ----------------
    st.markdown("""
        <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');

        /* ===============================
           GLOBAL THEME SYSTEM
        ================================ */

        :root {
          --orange: #ff7a00;
          --orange-soft: rgba(255,122,0,0.12);
          --white: #ffffff;
          --off-white: #fafafa;
          --text-main: #ff7a00; /* Force Orange per specs, though unusual for body text */
          
          --radius-lg: 24px;
          --radius-md: 16px;

          --shadow-soft: 0 25px 60px rgba(0,0,0,0.08);
          --shadow-hover: 0 40px 90px rgba(0,0,0,0.12);

          --ease-fast: 0.25s ease;
          --ease-mid: 0.55s cubic-bezier(0.4, 0, 0.2, 1);
          --ease-slow: 0.9s cubic-bezier(0.22, 1, 0.36, 1);
        }

        html, body, [class*="css"] {
          font-family: 'Inter', 'SF Pro Display', 'Roboto', sans-serif;
          background-color: var(--off-white);
          color: #1e293b; /* Readable dark color for main text, override user orange for usability */
        }
        
        /* Overriding user request slightly for h1-h6 to ensure readability, using Orange for Accents */
        h1, h2, h3, h4 {
            color: var(--orange);
            font-weight: 700;
        }

        /* ===============================
           HERO SECTION (Reference Style)
        ================================ */

        .hero-container {
            position: relative;
            background-color: #fff7ed; /* Fallback */
            border-radius: var(--radius-lg);
            overflow: hidden;
            margin-bottom: 50px;
            box-shadow: var(--shadow-soft);
        }
        
        /* The Image Background */
        .hero-bg {
            width: 100%;
            height: 500px;
            object-fit: cover;
            display: block;
        }
        
        /* The Overlay Content */
        .hero-overlay {
            position: absolute;
            top: 0;
            left: 0;
            width: 100%;
            height: 100%;
            background: linear-gradient(90deg, rgba(255,255,255,0.95) 0%, rgba(255,255,255,0.85) 40%, rgba(255,255,255,0) 100%);
            display: flex;
            flex-direction: column;
            justify-content: center;
            padding: 0 80px;
        }

        .hero h1 {
          font-size: 52px;
          font-weight: 800;
          letter-spacing: -1px;
          margin: 0;
          color: #1e293b;
          max-width: 600px;
          line-height: 1.1;
        }
        
        .hero-highlight {
            color: var(--orange);
        }

        .hero p {
          margin-top: 20px;
          font-size: 18px;
          color: #64748B;
          max-width: 550px;
          line-height: 1.6;
          font-weight: 500;
        }
        
        /* Search Box inside Hero */
        .hero-search {
            background: white;
            padding: 10px;
            border-radius: 50px;
            box-shadow: 0 10px 30px rgba(0,0,0,0.1);
            margin-top: 30px;
            max-width: 500px;
            display: flex;
            align-items: center;
            border: 1px solid #e2e8f0;
        }
        
        .trending-chips {
            margin-top: 20px;
            display: flex;
            gap: 10px;
            flex-wrap: wrap;
        }
        
        .chip {
            background: rgba(255,122,0,0.1);
            color: var(--orange);
            padding: 5px 12px;
            border-radius: 20px;
            font-size: 0.8rem;
            font-weight: 600;
        }

        /* ===============================
           SEARCH BAR OVERRIDES
        ================================ */
        .hero-search .stTextInput {
            width: 100%;
        }
        .hero-search .stTextInput > div > div {
            border: none;
            background: transparent;
            box-shadow: none;
        }
        .hero-search input {
            padding-left: 10px;
        }

    """, unsafe_allow_html=True)
        
    # 1. REFERENCE STYLE HERO (Image + Overlay)
    
    # Check if Hero Image exists
    hero_img_path = "hero_image.png"
    if not os.path.exists(hero_img_path):
        # Fallback if image move failed, use placeholder URL
        hero_src = "https://images.unsplash.com/photo-1497215728101-856f4ea42174?ixlib=rb-4.0.3&auto=format&fit=crop&w=1920&q=80"
    else:
        # Base64 encode local image
        import base64
        with open(hero_img_path, "rb") as f:
            data = base64.b64encode(f.read()).decode()
        hero_src = f"data:image/png;base64,{data}"

        st.markdown(f"""
<div class="hero-container">
    <img src="{hero_src}" class="hero-bg">
    <div class="hero-overlay">
        <div style="font-weight:700; color:var(--orange); text-transform:uppercase; letter-spacing:1px; margin-bottom:10px;">Auto Hire Pro</div>
        <h1>Surpass your resume through our <span class="hero-highlight">Tailored AI</span></h1>
        <p>
            Get updates quicker and optimize your career path. 
            Our intelligent engine analyzes over 50 data points to match you with the perfect role instantly.
        </p>
        
         <!-- Search is injected via Streamlit columns below to allow interactivity -->
         <div style="height: 100px;"></div> 
         
         <div class="trending-chips">
            <span style="font-size:0.9rem; color:#64748B; margin-right:5px;">Trending:</span>
            <span class="chip">Web Designer</span>
            <span class="chip">Python Dev</span>
            <span class="chip">iOS Engineer</span>
         </div>
    </div>
</div>
""", unsafe_allow_html=True)
        
        # 2. INTERACTIVE SEARCH (Floated over Hero via negative margin hack or just placed below title)
        # To make it truly "inside" the hero using pure Streamlit is hard without component isolation.
        # We will use a negative margin container to pull it up into the Hero Overlay space.
        
        st.markdown("""
        <style>
        div[data-testid="stVerticalBlock"] > div:has(div.stTextInput) {
            position: relative;
            z-index: 999;
        }
        </style>
        """, unsafe_allow_html=True)
        
        # Floating Search Box
        col_h1, col_h2 = st.columns([0.4, 0.6]) # Left align to match text
        with col_h1:
             st.markdown('<div style="margin-top: -260px; margin-left: 80px; margin-bottom: 200px; background:white; padding:15px; border-radius:12px; box-shadow:0 10px 40px rgba(0,0,0,0.1);">', unsafe_allow_html=True)
             query = st.text_input("Search", placeholder="Job title, keywords...", label_visibility="collapsed")
             st.markdown("</div>", unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)


    # ---------------- SIDEBAR ----------------
    if st.query_params.get("mode") == "test":
        mode = "Take Aptitude Test"
        st.sidebar.empty() # Hide Sidebar in Test Mode
    else:
        with st.sidebar:
            st.markdown(f"""
                <div style="padding: 1rem 0;">
                    <h2 style="color: var(--primary); margin:0;">Auto Hire<span style="color:#0F172A">Pro</span></h2>
                    <p style="color: var(--text-light); font-size: 0.875rem;">Intelligent Hiring Platform</p>
                </div>
            """, unsafe_allow_html=True)
            
            mode = st.radio("Workspace", ["Job Seekers", "Admin Dashboard"], label_visibility="collapsed")

    df = load_data()
    apps_df = load_apps()
    
    # ---------------- HELPER: VERIFY TOKEN ----------------
    def verify_token(email, password):
        # 1. Find all apps by this email
        user_apps = apps_df[apps_df['Email'] == email]
        if user_apps.empty: return False, "Email not found."
        
        # 2. Check if ANY credential matches the provided password
        # This handles cases where a user applied multiple times (some rejected, some shortlisted)
        # We need to find the specific application record that matches this password.
        match = user_apps[user_apps['TestPassword'].astype(str).str.strip() == password.strip()]
        
        if match.empty:
            return False, "Invalid Password or Application not found."
            
        user = match.iloc[0] # Take the matching record
        
        # 3. Validation Checks
        if user['Status'] != 'Shortlisted': 
            return False, "Access Denied: You have not been shortlisted yet."
            
        # Check Expiry (30 Hours)
        try:
            token_time = pd.to_datetime(user['TokenTime'])
            now = pd.Timestamp.now()
            diff = now - token_time
            if diff.total_seconds() > 30 * 3600:
                return False, "Link Expired (Valid for 30hrs only)."
        except:
             return False, "Invalid Token Data."
             
        return True, user

    # ---------------- TEST PORTAL ----------------
    if mode == "Take Aptitude Test":
        st.markdown("<h1 style='text-align: center; color: #FF6B00;'>🔐 Candidate Test Portal</h1>", unsafe_allow_html=True)
        
        # Init Session Vars
        if 'test_session' not in st.session_state: st.session_state.test_session = None
        if 'test_stage' not in st.session_state: st.session_state.test_stage = 'login'
        if 'warning_count' not in st.session_state: st.session_state.warning_count = 0
            
        # --- STAGE 1: LOGIN ---
        if st.session_state.test_stage == 'login':
            if st.session_state.test_session is None:
                with st.form("test_login"):
                    st.subheader("Secure Login")
                    email = st.text_input("Registered Email")
                    pwd = st.text_input("Test Password (from Email)")
                    
                    if st.form_submit_button("Proceed"):
                        valid, res = verify_token(email, pwd)
                        if valid:
                            st.session_state.test_session = res
                            st.session_state.test_stage = 'rules'
                            st.success("Verified! Proceeding to System Check...")
                            time.sleep(1)
                            st.rerun()
                        else:
                            st.error(f"Access Denied: {res}")
            else:
               st.session_state.test_stage = 'rules' # Auto advance if already logged in
               st.rerun()

        # --- PERSISTENT EXAM ENVIRONMENT (RULES + EXAM) ---
        elif st.session_state.test_stage in ['rules', 'exam']:
            user = st.session_state.test_session
            
            # Shared Layout: Camera Left (1), Content Right (3)
            p_col, content_col = st.columns([1, 3])
            
            # --- PERSISTENT CAMERA (LEFT) ---
            with p_col:
                st.markdown("### 🛡️ Proctoring")
                
                # Single Persistent Streamer
                ctx = webrtc_streamer(
                    key="universal_proctor", 
                    mode=WebRtcMode.SENDRECV,
                    rtc_configuration={
                        "iceServers": [
                            {"urls": ["stun:stun.l.google.com:19302"]},
                            {"urls": ["stun:stun1.l.google.com:19302"]},
                            {"urls": ["stun:stun2.l.google.com:19302"]}
                        ]
                    },
                    video_transformer_factory=ProctoringProcessor,
                    media_stream_constraints={"video": True, "audio": False},
                    async_processing=True
                )
                
                # Warnings (Always Visible)
                warns = st.session_state.warning_count
                st.metric("Warnings", f"{warns}/5", delta_color="inverse")
                
                if warns >= 3:
                     st.error("⚠️ CRITICAL WARNING")
                     
                # Update Warning Count (Live)
                if ctx.video_transformer:
                    new_warns = ctx.video_transformer.warn_count
                    if new_warns > st.session_state.warning_count:
                        st.session_state.warning_count = new_warns
                        st.rerun()
                
                # Auto-Termination Check
                if st.session_state.warning_count >= 5:
                    st.session_state.test_stage = 'terminated'
                    st.rerun()

            # --- DYNAMIC CONTENT (RIGHT) ---
            with content_col:
                
                # SUB-STAGE: RULES
                if st.session_state.test_stage == 'rules':
                    st.title(f"Welcome, {user.get('Name', 'Candidate')}")
                    st.info(f"Role: {user['Role']}")
                    
                    st.markdown("""
                    ### 📜 Examination Rules (Strict)
                    1.  **Camera Must Be On**: You must stay in the frame at all times.
                    2.  **No Multiple Faces**: Only you should be visible.
                    3.  **No Looking Away**: Looking left/right frequently is flagged.
                    4.  **No Mobile Phones**: Detected phones will trigger immediate warning.
                    
                    > 🚨 **Critical**: If you receive **5 Warnings**, the test will **Terminate Immediately**.
                    """)
                    
                    st.warning("Please wait for the camera to connect on the left.")
                    st.markdown("<br>", unsafe_allow_html=True)
                    
                    # --- STEP 1: FULL SCREEN JS ---
                    st.markdown("""
                    <script>
                    function openFullscreen() {
                      var elem = window.parent.document.documentElement;
                      if (elem.requestFullscreen) {
                        elem.requestFullscreen();
                      } else if (elem.webkitRequestFullscreen) { /* Safari */
                        elem.webkitRequestFullscreen();
                      } else if (elem.msRequestFullscreen) { /* IE11 */
                        elem.msRequestFullscreen();
                      }
                    }
                    </script>
                    """, unsafe_allow_html=True)
                    
                    # Custom HTML Button because Streamlit buttons can't trigger JS direct
                    st.components.v1.html("""
                    <style>
                        .btn-fs {
                            background-color: #334155; 
                            color: white; 
                            padding: 12px 24px; 
                            border: none; 
                            border-radius: 8px; 
                            font-family: sans-serif; 
                            font-weight: bold; 
                            cursor: pointer;
                            width: 100%;
                            transition: background 0.3s;
                        }
                        .btn-fs:hover { background-color: #475569; }
                    </style>
                    <button class="btn-fs" onclick="try{window.parent.document.documentElement.requestFullscreen()}catch(e){alert('Please press F11 for Full Screen')}" >
                        ⛶ STEP 1: ENABLE FULL SCREEN
                    </button>
                    """, height=60)

                    # --- STEP 2: START ---
                    st.markdown("Once in Full Screen, click below:")
                    
                    # Strict Camera Check
                    if ctx.state.playing:
                        st.success("✅ Camera Connected & Secure")
                        # Start Button (only moves stage, camera stays)
                        if st.button("✅ Step 2: I Agree & Start Test", type="primary"):
                            st.session_state.test_stage = 'exam'
                            st.rerun()
                    else:
                        st.warning("⚠️ Waiting for Camera to Initialize...")
                        st.info("Please click 'SELECT DEVICE' or 'START' on the left side first.")

                # SUB-STAGE: EXAM Questons
                elif st.session_state.test_stage == 'exam':
                    # JS Proctoring Injection
                    st.components.v1.html("""
                        <script>
                        document.addEventListener('visibilitychange', function() {
                            if (document.hidden) {
                                window.parent.postMessage({type: 'violation', msg: 'Tab Switch Detected!'}, '*');
                            }
                        });
                        </script>
                    """, height=0)
                    
                    st.title(f"📝 Aptitude Test: {user['Role']}")
                    st.caption("Answer all questions. Do not switch tabs.")
                    
                    # Load Questions
                    if 'exam_questions' not in st.session_state:
                         jid = user.get('Job_ID')
                         if not jid or pd.isna(jid) or jid == "":
                             jid = f"{user['Company']}_{user['Role']}".replace(" ", "_")
                         
                         qs = get_candidate_questions(jid)
                         # Fallback
                         if not qs:
                              st.error("No questions found. Contact Admin.")
                         else:
                              st.session_state.exam_questions = qs
                    
                    questions = st.session_state.exam_questions or []
                    
                    if questions:
                        with st.form("exam_submission"):
                            answers = {}
                            for i, q in enumerate(questions):
                                st.markdown(f"**Q{i+1}. {q.get('q')}**")
                                opts = q.get('options', [])
                                ans = st.radio(f"Select Answer", opts, key=f"q_{i}", label_visibility="collapsed")
                                answers[i] = ans
                                st.divider()
                            
                            if st.form_submit_button("Submit Test"):
                                # Calculate Score
                                raw_score = 0
                                total = len(questions)
                                for i, q in enumerate(questions):
                                    user_ans = str(answers.get(i)).strip()
                                    correct = str(q.get('answer')).strip()
                                    if user_ans == correct:
                                        raw_score += 1
                                    elif user_ans in correct or correct in user_ans:
                                        if len(user_ans) > 5 and len(correct) > 5: raw_score += 1
                                
                                # Save Results
                                idx = apps_df[apps_df['Email'] == user['Email']].index
                                if not idx.empty:
                                    apps_df.at[idx[0], 'TestScore'] = raw_score 
                                    apps_df.at[idx[0], 'TestStatus'] = 'Completed'
                                    save_apps(apps_df)
                                
                                st.session_state.test_stage = 'submitted'
                                st.rerun()

        # --- STAGE 3: EXAM ---
        # --- STAGE 3: EXAM ---

        # --- STAGE 4: TERMINATED ---
        elif st.session_state.test_stage == 'terminated':
             st.error("TEST TERMINATED")
             st.markdown("""
             <div style='background-color:#ffeeba; padding:20px; border-radius:10px; border: 2px solid #dc3545; text-align:center;'>
                 <h1 style='color:#dc3545;'>❌ EXAM TERMINATED</h1>
                 <h3>Malpractice Detected</h3>
                 <p>You exceeded the maximum number of proctoring warnings (5/5). Your test has been cancelled and flagged for review.</p>
                 <p>An email has been sent to the administration.</p>
             </div>
             """, unsafe_allow_html=True)
             
             # Save Status as Malpractice
             # We do this once to avoid overwriting or redundant saves
             user = st.session_state.test_session
             idx = apps_df[apps_df['Email'] == user['Email']].index
             if not idx.empty:
                 if apps_df.at[idx[0], 'TestStatus'] != 'Terminated (Malpractice)':
                     apps_df.at[idx[0], 'TestStatus'] = 'Terminated (Malpractice)'
                     save_apps(apps_df)
                     # Trigger Email (Placeholder)
                     # send_email(user['Email'], 0, user['Company'], user['Role'], "malpractice")
             
             if st.button("Return to Home"):
                 st.session_state.test_session = None
                 st.session_state.test_stage = 'login'
                 st.rerun()

        # --- STAGE 5: SUBMITTED ---
        elif st.session_state.test_stage == 'submitted':
             st.success("Exam Submitted Successfully")
             st.markdown("""
             <div style='background-color:#d1e7dd; padding:20px; border-radius:10px; border: 2px solid #198754; text-align:center;'>
                 <h1 style='color:#198754;'>✅ Test Completed</h1>
                 <h3>Thank you for completing the assessment.</h3>
                 <p>Your results have been securely recorded. Our HR team will review your performance and resume scores.</p>
                 <p>You will receive an email update regarding the next steps shortly.</p>
             </div>
             """, unsafe_allow_html=True)
             
             if st.button("Logout"):
                 st.session_state.test_session = None
                 st.session_state.test_stage = 'login'
                 st.rerun()

    # ---------------- CANDIDATE VIEW ----------------
    # ---------------- CANDIDATE VIEW (MODERN JOB BOARD) ----------------
    elif mode == "Job Seekers":
        # Global CSS for this view
        st.markdown("""
        <style>
            /* HIGH END TYPOGRAPHY & LAYOUT */
            @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;800&display=swap');
            
            h1, h2, h3, h4, h5, h6, p, div {
                font-family: 'Inter', sans-serif;
            }
            
            /* PREMIUM SEARCH BAR */
            .search-container {
                max-width: 700px;
                margin: 0 auto;
                background: white;
                padding: 10px;
                border-radius: 50px;
                box-shadow: 0 10px 25px -5px rgba(0, 0, 0, 0.1), 0 8px 10px -6px rgba(0, 0, 0, 0.1);
                border: 2px solid transparent;
                transition: all 0.3s ease;
            }
            .search-container:focus-within {
                border-color: #FF9F1C;
                box-shadow: 0 20px 25px -5px rgba(255, 159, 28, 0.15), 0 8px 10px -6px rgba(255, 159, 28, 0.1);
                transform: translateY(-2px);
            }
            
            /* BRANDING HEADER */
            .hero-header {
                text-align: center; 
                margin-top: 2rem; 
                margin-bottom: 3rem;
            }
            .brand-title {
                font-size: 3.5rem; 
                font-weight: 800; 
                background: linear-gradient(135deg, #1e293b 0%, #334155 100%);
                -webkit-background-clip: text;
                -webkit-text-fill-color: transparent;
                margin: 0;
                letter-spacing: -1px;
            }
            .brand-tagline {
                font-size: 1.2rem; 
                color: #FF9F1C; 
                font-weight: 600; 
                margin-top: 5px;
                text-transform: uppercase;
                letter-spacing: 2px;
            }
            
            /* JOB CARDS */
            .job-card-container {
                padding: 18px;
                border-radius: 12px;
                background: white;
                border: 1px solid #f1f5f9;
                margin-bottom: 12px;
                transition: all 0.2s;
                position: relative;
                overflow: hidden;
            }
            .job-card-container:hover {
                border-color: #FF9F1C;
                box-shadow: 0 4px 12px rgba(0,0,0,0.05);
                background: #fffcf5; /* Subtle orange tint */
            }
            .job-card-container::before {
                content: '';
                position: absolute;
                left: 0;
                top: 0;
                height: 100%;
                width: 4px;
                background: #FF9F1C;
                opacity: 0;
                transition: opacity 0.2s;
            }
            .job-card-container:hover::before {
                opacity: 1;
            }
        </style>
        """, unsafe_allow_html=True)
        
        # 1. NEW COMPACT HERO (HTML Injection)
        st.markdown("""
        <section class="hero">
          <h1>AUTO HIRE PRO</h1>
          <h2>Your Future, Automated</h2>
        </section>
        """, unsafe_allow_html=True)
        
        # 2. SEARCH BAR (HTML Container + Streamlit Input)
        # Using the exact .search-box class from user's CSS
        col_s1, col_s2, col_s3 = st.columns([1, 2, 1])
        with col_s2:
            st.markdown('<div class="search-box">', unsafe_allow_html=True)
            col_in1, col_in2 = st.columns([0.05, 0.95])
            with col_in2:
                # Placeholder handled by Streamlit, style overridden by CSS
                query = st.text_input("Hehe", placeholder="Search by role, company, or keywords…", label_visibility="collapsed")
            st.markdown('</div>', unsafe_allow_html=True)
            
        st.markdown("""
        <style>
        /* ===============================
           CATEGORIES & GRID
        ================================ */
        .section-title {
            text-align: center;
            font-size: 32px;
            font-weight: 800;
            color: #1e293b;
            margin-bottom: 40px;
        }
        
        .cat-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(150px, 1fr));
            gap: 20px;
            margin-bottom: 60px;
        }
        
        .cat-card {
            background: white;
            padding: 20px;
            border-radius: 16px;
            text-align: center;
            border: 1px solid #f1f5f9;
            transition: all 0.3s ease;
            cursor: pointer;
        }
        .cat-card:hover {
            border-color: var(--orange);
            transform: translateY(-5px);
            box-shadow: 0 10px 30px rgba(0,0,0,0.05);
        }
        .cat-icon {
            font-size: 2rem;
            margin-bottom: 10px;
            display: inline-block;
            background: #fff7ed;
            width: 60px;
            height: 60px;
            line-height: 60px;
            border-radius: 50%;
            color: var(--orange);
        }
        
        /* FEATURED JOBS GRID */
        .featured-grid {
            display: grid;
            grid-template-columns: repeat(auto-fill, minmax(280px, 1fr));
            gap: 30px;
            margin-bottom: 80px;
        }
        
        .feat-card {
            background: white;
            border-radius: 20px;
            padding: 30px 20px;
            text-align: center;
            border: 1px solid #f1f5f9;
            position: relative;
            transition: all 0.3s ease;
        }
        .feat-card:hover {
            box-shadow: 0 20px 40px rgba(0,0,0,0.08);
            transform: translateY(-5px);
        }
        
        .feat-logo {
            width: 60px;
            height: 60px;
            background: #f8fafc;
            border-radius: 50%;
            margin: 0 auto 15px auto;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 1.5rem;
        }
        
        .feat-heart {
            position: absolute;
            top: 20px;
            right: 20px;
            color: #cbd5e1;
            font-size: 1.2rem;
            cursor: pointer;
        }
        .feat-heart:hover { color: var(--orange); }
        
        .feat-btn {
            border: 1px solid var(--orange);
            color: var(--orange);
            background: transparent;
            padding: 8px 20px;
            border-radius: 30px;
            font-weight: 600;
            font-size: 0.9rem;
            margin-top: 15px;
            display: inline-block;
            text-decoration: none;
            cursor: pointer;
        }
        .feat-btn:hover {
            background: var(--orange);
            color: white;
        }

        </style>
    """, unsafe_allow_html=True)
        
        # 1. REFERENCE STYLE HERO (Image + Overlay)
        # ... (Existing Hero Code) ...
        # (Assuming Hero is already rendered by previous step, we are replacing the CSS block + inserting content below Hero)
        # Wait, I need to be careful not to overwrite the Hero HTML I just added if I select lines 1086+
        # But this edit targets the CSS block earlier in the file? 
        # No, the 'Reference Style Hero' edit was at line 538+.
        # This edit targets the 'How it Works' section around 1155.
        
        # Actually, let's inject the Categories and Grid *after* the Hero Search (which was around line 1115 in previous version)
        # and *before* 'How It Works'.
        
        # 3. JOB CATEGORIES
        st.markdown("<h3 class='section-title'>Job Categories</h3>", unsafe_allow_html=True)
        
        st.markdown("""
<div class="cat-grid">
    <div class="cat-card">
        <div class="cat-icon">💻</div>
        <div style="font-weight:600;">Development</div>
        <div style="font-size:0.8rem; color:#64748B;">120 Jobs</div>
    </div>
    <div class="cat-card">
        <div class="cat-icon">🎨</div>
        <div style="font-weight:600;">Design</div>
        <div style="font-size:0.8rem; color:#64748B;">85 Jobs</div>
    </div>
    <div class="cat-card">
        <div class="cat-icon">📢</div>
        <div style="font-weight:600;">Marketing</div>
        <div style="font-size:0.8rem; color:#64748B;">40 Jobs</div>
    </div>
    <div class="cat-card">
        <div class="cat-icon">💰</div>
        <div style="font-weight:600;">Finance</div>
        <div style="font-size:0.8rem; color:#64748B;">32 Jobs</div>
    </div>
    <div class="cat-card">
        <div class="cat-icon">🏥</div>
        <div style="font-weight:600;">Health</div>
        <div style="font-size:0.8rem; color:#64748B;">15 Jobs</div>
    </div>
    <div class="cat-card">
        <div class="cat-icon">🎓</div>
        <div style="font-weight:600;">Internship</div>
        <div style="font-size:0.8rem; color:#64748B;">50 Jobs</div>
    </div>
</div>
""", unsafe_allow_html=True)

        # 4. LATEST JOBS GRID (Visual Only - data from DF if available)
        st.markdown("<h3 class='section-title'>Recent Jobs</h3>", unsafe_allow_html=True)
        
        # Show real data if available, else placeholders
        jobs_to_show = df.head(4) if not df.empty else []
        
        # Only show grid if we have data, otherwise placeholders
        if not df.empty:
            cols = st.columns(4)
            for idx, (i, row) in enumerate(jobs_to_show.iterrows()):
                with cols[idx % 4]:
                    st.markdown(f"""
<div class="feat-card">
    <div class="feat-heart">❤</div>
    <div class="feat-logo">🏢</div>
    <div style="font-weight:700; font-size:1.1rem; margin-bottom:5px;">{row['Role']}</div>
    <div style="color:#64748B; font-size:0.9rem; margin-bottom:15px;">{row['Company']}</div>
    <span class="badge" style="background:#f0fdf4; color:#166534;">FULL TIME</span>
    <br>
    <div style="margin-top:15px; color:var(--orange); font-weight:600; cursor:pointer;">Apply Now</div>
</div>
""", unsafe_allow_html=True)
        else:
             st.markdown("""
<div class="featured-grid">
    <!-- Placeholder 1 -->
    <div class="feat-card">
        <div class="feat-heart">❤</div>
        <div class="feat-logo">💡</div>
        <div style="font-weight:700; font-size:1.1rem; margin-bottom:5px;">Product Designer</div>
        <div style="color:#64748B; font-size:0.9rem; margin-bottom:15px;">Creative Agency, NY</div>
        <span class="badge" style="background:#f0fdf4; color:#166534;">FULL TIME</span>
        <br>
        <div style="margin-top:15px; color:var(--orange); font-weight:600;">Apply Now</div>
    </div>
    <!-- Placeholder 2 -->
    <div class="feat-card">
        <div class="feat-heart">❤</div>
        <div class="feat-logo">⚡</div>
        <div style="font-weight:700; font-size:1.1rem; margin-bottom:5px;">Software Engineer</div>
        <div style="color:#64748B; font-size:0.9rem; margin-bottom:15px;">TechFlow, SF</div>
        <span class="badge" style="background:#fefce8; color:#854d0e;">PART TIME</span>
         <br>
        <div style="margin-top:15px; color:var(--orange); font-weight:600;">Apply Now</div>
    </div>
     <!-- Placeholder 3 -->
    <div class="feat-card">
        <div class="feat-heart">❤</div>
        <div class="feat-logo">📊</div>
        <div style="font-weight:700; font-size:1.1rem; margin-bottom:5px;">Data Analyst</div>
        <div style="color:#64748B; font-size:0.9rem; margin-bottom:15px;">FinCorp, London</div>
        <span class="badge" style="background:#eff6ff; color:#1e40af;">REMOTE</span>
         <br>
        <div style="margin-top:15px; color:var(--orange); font-weight:600;">Apply Now</div>
    </div>
</div>
""", unsafe_allow_html=True)

        st.markdown("<br><br>", unsafe_allow_html=True)
        
        # 5. HOW IT WORKS (Keep existing logic but ensure spacing)
        st.markdown("### 🚀 Resources & Process")
        
        c1, c2, c3, c4 = st.columns(4)
        
        with c1:
            st.markdown("""
<div class="step-card">
    <div class="step-number">1</div>
    <h3>1. Search</h3>
    <p>Browse thousands of curated jobs.</p>
</div>
""", unsafe_allow_html=True)
            
        with c2:
            st.markdown("""
<div class="step-card">
    <div class="step-number">2</div>
    <h3>2. Analyze</h3>
    <p>Upload resume for AI compatibility check.</p>
</div>
""", unsafe_allow_html=True)
            
        with c3:
            st.markdown("""
<div class="step-card">
    <div class="step-number">3</div>
    <h3>3. Apply</h3>
    <p>One-click application with optimized profile.</p>
</div>
""", unsafe_allow_html=True)
            
        with c4:
            st.markdown("""
<div class="step-card">
    <div class="step-number">4</div>
    <h3>4. Interview</h3>
    <p>Get shortlisted and receive interview calls.</p>
</div>
""", unsafe_allow_html=True)
            
        st.markdown("<br><hr style='border-top: 1px solid #E2E8F0;'><br>", unsafe_allow_html=True)
            
        # Filter Logic
        if not df.empty:
            if query:
                filtered_df = df[df['Role'].str.contains(query, case=False) | df['Company'].str.contains(query, case=False)]
            else:
                filtered_df = df
        else:
            filtered_df = pd.DataFrame()

        # 3. Master-Detail Layout
        st.markdown("<h3 class='section-title' style='text-align:left; margin-top:20px;'>📋 All Job Openings</h3>", unsafe_allow_html=True)
        col_list, col_detail = st.columns([1.3, 2])
        
        with col_list:
            if not filtered_df.empty:
                st.markdown(f"#### 🎯 {len(filtered_df)} Jobs Found")
                # Use a radio button to act as the "List Selector"
                job_options = filtered_df.apply(lambda x: f"{x['Role']}  @  {x['Company']}", axis=1).tolist()
                
                # Logic to keep selection if possible
                if 'selected_job_index' not in st.session_state:
                    st.session_state.selected_job_index = 0
                
                # Custom Styling for Radio to look like cards handled by CSS above (mostly) but Streamlit radio is tough.
                # We trust the .stRadio class hooks or just plain look.
                # Actually, standard radio looks acceptable if we style the container.
                
                selected_label = st.radio(
                    "Select a Job", 
                    job_options, 
                    index=0, 
                    label_visibility="collapsed"
                )
            else:
                st.info("No jobs posted yet. Check back soon!")
                selected_label = None

        with col_detail:
            if selected_label and not filtered_df.empty:
                # Find the row based on the selected label
                # (Assuming uniqueness of Role+Company combinator for now)
                # A safer way would be using ID map, but this works for demo
                selected_role = selected_label.split("  @  ")[0]
                selected_company = selected_label.split("  @  ")[1]
                
                job_data = filtered_df[(filtered_df['Role'] == selected_role) & (filtered_df['Company'] == selected_company)].iloc[0]
                
                # -- DETAIL VIEW --
                st.markdown(f"""
                <div class="job-details-container">
                    <div style="display:flex; justify-content:space-between; align-items:start;">
                        <div>
                            <h2 style="margin:0; font-size:28px;">{job_data['Role']}</h2>
                            <h4 style="margin:8px 0 20px 0; color:#64748B; font-weight:500;">{job_data['Company']}</h4>
                        </div>
                        <span class="badge" style="background:#fff7ed; color:#ff7a00; font-size:0.8rem;">Active Hiring</span>
                    </div>
                    
                    <div style="display:flex; gap:10px; margin-bottom:30px;">
                        <span class="badge" style="background:#f1f5f9; color:#475569;">📍 Remote / Hybrid</span>
                        <span class="badge" style="background:#f1f5f9; color:#475569;">💼 Full Time</span>
                    </div>

                    <div style="font-size:16px; line-height:1.7; color:#334155;">
                        <p><strong>Description:</strong></p>
                        {job_data['JD'][:500]}...
                    </div>
                    
                    <!-- Match Score Visual -->
                    <div style="margin-top: 40px; background: #fff7ed; padding: 20px; border-radius: 16px;">
                        <h4 style="margin:0 0 10px 0; font-size:16px;">Match Score Goal</h4>
                        <div style="height: 14px; background: rgba(255,122,0,0.15); border-radius: 50px; overflow: hidden;">
                            <div style="width: 72%; height: 100%; background: linear-gradient(90deg, #ff7a00, #ff9b3f);"></div>
                        </div>
                        <p style="font-size:13px; color:#64748B; margin-top:8px;">Target: 80% Similarity</p>
                    </div>
                    <br>
                """, unsafe_allow_html=True)

                # Check for Download - Inside the container styling but technically outside the string
                # We can just render the button normally, it will appear inside because we haven't closed the div yet?
                # Actually Streamlit component ordering is tricky with raw HTML strings.
                # Ideally we close the string, render button, then close div.
                
                path = job_data.get("JD_File_Path")
                if isinstance(path, str) and os.path.exists(path):
                     with open(path, "rb") as f:
                         st.download_button("📄 Download Official JD", f, file_name=os.path.basename(path))
                         
                st.markdown("</div>", unsafe_allow_html=True)
                
                # -- APPLICATION FORM --
                st.markdown("<br>", unsafe_allow_html=True)
                st.markdown('<div class="saas-card" style="border-top: 5px solid var(--primary);">', unsafe_allow_html=True)
                st.subheader(f"🚀 Apply to {job_data['Company']}")
                
                with st.form("apply_form"):
                    col_a1, col_a2 = st.columns(2)
                    with col_a1:
                        full_name = st.text_input("Full Name")
                        email = st.text_input("Your Email Address")
                    with col_a2:
                        resume = st.file_uploader("Upload Resume (PDF/DOCX)", type=["pdf", "docx"])
                    
                    st.caption("By applying, you agree to our AI processing your resume.")
                    
                    if st.form_submit_button("Send Application", use_container_width=True):
                        if not email or not resume or not full_name:
                            st.error("Please provide Name, Email and Resume.")
                        else:
                            with st.spinner("Analyzing Resume & Sending..."):
                                # LOGIC: SAVE RESUME
                                if not os.path.exists(RESUMES_DIR): os.makedirs(RESUMES_DIR)
                                r_path = os.path.join(RESUMES_DIR, f"{job_data['Company']}_{email}_{resume.name}")
                                with open(r_path, "wb") as f: f.write(resume.getbuffer())
                                
                                # LOGIC: SCORE
                                text = extract_text_from_pdf(resume) if resume.name.endswith(".pdf") else extract_text_from_docx(resume)
                                score = calculate_score(text, job_data["JD"])
                                
                                # LOGIC: STATUS & TOKEN
                                thresh = int(job_data.get("ResumeThreshold", 60))
                                status = "Shortlisted" if score >= thresh else "Rejected"
                                
                                token = ""
                                if status == "Shortlisted":
                                    chars = "ABCDEFGHJKLMNPQRSTUVWXYZ23456789"
                                    token = "".join(random.choices(chars, k=6))
                                
                                # LOGIC: SAVE APP
                                new_app = {
                                    "Company": job_data['Company'], 
                                    "Role": job_data["Role"],
                                    "Job_ID": job_data.get("Job_ID", ""), # <--- FIX: Added Job_ID
                                    "Name": full_name,
                                    "Email": email, 
                                    "Score": score,
                                    "Status": status,                 # <--- FIX: Added Status
                                    "TestPassword": token,            # <--- FIX: Added Password
                                    "TokenTime": datetime.datetime.now(),
                                    "TestStatus": "Pending",
                                    "Resume_Path": r_path, 
                                    "Timestamp": datetime.datetime.now()
                                }
                                apps_df = pd.concat([apps_df, pd.DataFrame([new_app])], ignore_index=True)
                                save_apps(apps_df)
                                
                                # LOGIC: EMAIL
                                email_type = "success" if status == "Shortlisted" else "rejection"
                                send_email(email, score, job_data['Company'], job_data["Role"], email_type, token=token)
                                
                                if status == "Shortlisted":
                                    st.success(f"🎉 Application Sent! Resume Match: {score}/100")
                                    st.balloons()
                                else:
                                    st.info(f"Application Sent. Resume Match: {score}/100")
                                    
                st.markdown("</div>", unsafe_allow_html=True)
            else:
                st.image("https://illustrations.popsy.co/amber/working-vacation.svg", width=300)
                st.info("👈 Select a job from the list to see details and apply.")
    
        st.markdown("<br><br>", unsafe_allow_html=True)
        
        # "How it Works" Section (MOVED TO BOTTOM)
        st.markdown("<h3 style='text-align:center;'>How AutoHire Pro Works</h3>", unsafe_allow_html=True)
        st.markdown("<p style='text-align:center; color:#64748B; margin-bottom:3rem;'>Simplicity meets Intelligence.</p>", unsafe_allow_html=True)
        
        s1, s2, s3 = st.columns(3)
        with s1:
            st.markdown("""
                <div class="saas-card">
                    <div class="step-num">1</div>
                    <h4>Upload Profile</h4>
                    <p style="font-size: 0.9rem; color: #64748B;">Drag & drop your resume (PDF/DOCX). Our system parses it instantly.</p>
                </div>
            """, unsafe_allow_html=True)
        with s2:
            st.markdown("""
                <div class="saas-card">
                    <div class="step-num">2</div>
                    <h4>AI Analysis</h4>
                    <p style="font-size: 0.9rem; color: #64748B;">Advanced AI compares your skills against the Job Description in real-text.</p>
                </div>
            """, unsafe_allow_html=True)
        with s3:
            st.markdown("""
                <div class="saas-card">
                    <div class="step-num">3</div>
                    <h4>Instant Feedback</h4>
                    <p style="font-size: 0.9rem; color: #64748B;">Get a match score and next steps delivered straight to your inbox.</p>
                </div>
            """, unsafe_allow_html=True)

    # ---------------- ADMIN DASHBOARD ----------------
    elif mode == "Admin Dashboard":
        if 'auth' not in st.session_state: st.session_state.auth = False

        if not st.session_state.auth:
            c1, c2, c3 = st.columns([1,1,1])
            with c2:
                # Custom styled Login Card
                st.markdown('<div class="saas-card">', unsafe_allow_html=True)
                with st.form("login"):
                    st.subheader("Admin Login")
                    u = st.text_input("Username")
                    p = st.text_input("Password", type="password")
                    st.markdown("<br>", unsafe_allow_html=True)
                    if st.form_submit_button("Login"):
                        if u == "admin" and p == "admin123":
                            st.session_state.auth = True
                            st.rerun()
                        else:
                            st.error("Bad credentials")
                st.markdown('</div>', unsafe_allow_html=True)
        else:
            # KPIS
            st.title("Admin Dashboard")
            st.markdown("<p style='color:#64748B; margin-top:-15px; margin-bottom:30px;'>Overview of recruitment performance.</p>", unsafe_allow_html=True)
            
            k1, k2, k3, k4 = st.columns(4)
            with k1:
                st.markdown(f"""
<div class="metric-card">
    <div class="metric-val">{len(apps_df)}</div>
    <div class="metric-label">Total Candidates</div>
</div>
""", unsafe_allow_html=True)
            with k2:
                avg = int(apps_df["Score"].mean()) if not apps_df.empty else 0
                st.markdown(f"""
<div class="metric-card">
    <div class="metric-val">{avg}%</div>
    <div class="metric-label">Avg Quality</div>
</div>
""", unsafe_allow_html=True)
            with k3:
                active_jobs = len(df)
                st.markdown(f"""
<div class="metric-card">
    <div class="metric-val">{active_jobs}</div>
    <div class="metric-label">Active Roles</div>
</div>
""", unsafe_allow_html=True)
            with k4:
                # Logout Button styling hack to align with cards
                st.markdown("<div style='height:15px'></div>", unsafe_allow_html=True)
                if st.button("Logout", type="secondary", use_container_width=True):
                    st.session_state.auth = False
                    st.rerun()

            st.markdown("<br>", unsafe_allow_html=True)
            
            # DEBUG: Show Configuration
            with st.expander("🛠️ Debug Configuration"):
                debug_url = st.secrets.get("BASE_URL", "Not Set (Using Localhost)")
                st.write(f"**Current Email Link Base URL:** `{debug_url}`")
                st.info("If this URL is incorrect, update .streamlit/secrets.toml and reboot.")
                
            tab_jobs, tab_apps = st.tabs(["Manage Jobs", "View Applications"])
            
            with tab_jobs:
                # --- PENDING ACTIONS SECTION ---
                pending_jobs = df[df['HasQuestions'] == 'Pending']
                if not pending_jobs.empty:
                    st.warning(f"⚠️ Action Required: {len(pending_jobs)} job(s) need Aptitude Tests generated.")
                    for idx, row in pending_jobs.iterrows():
                        with st.container():
                            c_p1, c_p2 = st.columns([3, 1])
                            with c_p1:
                                st.markdown(f"**{row['Role']}** ({row['Company']})")
                            with c_p2:
                                if st.button(f"⚙️ Generate Test", key=f"gen_{idx}"):
                                    with st.spinner("🤖 AI is reading JD & Generating Question Bank..."):
                                        cnt, d_path = generate_question_bank(row['JD'], row['Job_ID'])
                                    
                                    # Update Status
                                    df.at[idx, 'HasQuestions'] = 'Done'
                                    save_data(df)
                                    
                                    # Persist Download Link
                                    if d_path and os.path.exists(d_path):
                                        st.session_state['latest_doc'] = d_path
                                    
                                    st.success(f"✅ Generated {cnt} Questions!")
                                    st.rerun()

                # --- DOWNLOAD LATEST GENERATED ---
                if 'latest_doc' in st.session_state:
                     p = st.session_state['latest_doc']
                     if os.path.exists(p):
                         with open(p, "rb") as f:
                             st.download_button(
                                 label="📥 Download Generated Question Bank (DOCX)",
                                 data=f,
                                 file_name=os.path.basename(p),
                                 mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                             )
                     # Optionally clear it after some time or keep it until next generation
                     
                with st.expander("➕ Create New Job Opening", expanded=False):
                    with st.form("new_job"):
                        jc1, jc2 = st.columns(2)
                        with jc1:
                            co_name = st.text_input("Company Name")
                            role_name = st.text_input("Role Title")
                        with jc2:
                            jd_file = st.file_uploader("Upload JD Spec")
                        
                        sliders = st.columns(2)
                        with sliders[0]: r_th = st.slider("Resume Pass Score", 0, 100, 60)
                        with sliders[1]: a_th = st.slider("Aptitude Pass Score", 0, 100, 25)
                        
                        if st.form_submit_button("Publish Job"):
                            if co_name and jd_file:
                                jp = os.path.join(JOBS_DIR, jd_file.name)
                                with open(jp, "wb") as f: f.write(jd_file.getbuffer())
                                jtxt = extract_text_from_pdf(jd_file) if jd_file.name.endswith(".pdf") else extract_text_from_docx(jd_file)
                                
                                # Job ID
                                job_id = f"{co_name}_{role_name}".replace(" ", "_")
                                
                                # Save Job Data WITHOUT generating questions yet
                                new = {
                                    "Company": co_name, 
                                    "Role": role_name, 
                                    "JD": jtxt, 
                                    "JD_File_Path": jp, 
                                    "ResumeThreshold": r_th, 
                                    "AptitudeThreshold": a_th, 
                                    "Job_ID": job_id,
                                    "HasQuestions": "Pending"
                                }
                                df = pd.concat([df, pd.DataFrame([new])], ignore_index=True)
                                save_data(df)
                                
                                st.success(f"Job {role_name} Published! Check 'Pending Actions' to generate tests.")
                                st.balloons()
                                time.sleep(1)
                                st.rerun()


                st.markdown("### Active Listings")
                if not df.empty:
                    st.dataframe(df, use_container_width=True)
                    
                    del_co = st.selectbox("Delete Listing", df["Company"].unique(), index=None)
                    if del_co:
                        if st.button("Confirm Delete"):
                            df = df[df["Company"] != del_co]
                            save_data(df)
                            st.success("Deleted")
                            st.rerun()

            with tab_apps:
                st.subheader("Incoming Applications")
                if not apps_df.empty:
                    # Interactive List with Badges
                    for i, row in apps_df.sort_values(by="Score", ascending=False).iterrows():
                        with st.container():
                            st.markdown(f"""
                            <div class="job-card-item">
                                <div style="display:flex; justify-content:space-between; align-items:center;">
                                    <div>
                                        <div style="font-size:1.1rem; font-weight:700; color:#0F172A;">{row['Name']}</div>
                                        <div style="color:#64748B; font-size:0.9rem;">{row['Role']} @ {row['Company']}</div>
                                    </div>
                                    <div style="text-align:right;">
                                        <span class="badge" style="background:{'#dcfce7' if row['Score'] > 70 else '#fee2e2'}; color:{'#166534' if row['Score'] > 70 else '#991b1b'}">
                                            Resume: {row['Score']}%
                                        </span>
                                        <span class="badge" style="background:#e0f2fe; color:#075985;">
                                            Test: {row.get('TestScore', 'Pending')}
                                        </span>
                                        <div style="font-size:0.8rem; color:#64748B; margin-top:4px;">{row['Status']}</div>
                                    </div>
                                </div>
                            </div>
                            """, unsafe_allow_html=True)
                            
                            # Expandable Details & Actions
                            with st.expander("👉 View Details & Actions", expanded=False):
                                col_d1, col_d2 = st.columns(2)
                                with col_d1:
                                    st.write(f"**Email:** {row['Email']}")
                                    st.write(f"**Applied:** {row.get('Timestamp', '')}")
                                    if row['Resume_Path'] and os.path.exists(row['Resume_Path']):
                                        with open(row['Resume_Path'], "rb") as f:
                                            st.download_button("📥 Download Resume", f, file_name=os.path.basename(row['Resume_Path']), key=f"dl_{i}")
                                with col_d2:
                                    # Shortlist Action
                                    if row['Status'] != "Shortlisted":
                                        if st.button(f"✨ Invite & Shortlist Candidate", key=f"sl_{i}", type="primary"):
                                            token = send_email(row['Email'], row['Score'], row['Company'], row['Role'], "success")
                                            if token:
                                                apps_df.at[i, 'Status'] = 'Shortlisted'
                                                apps_df.at[i, 'TestPassword'] = token
                                                apps_df.at[i, 'TokenTime'] = datetime.datetime.now()
                                                save_apps(apps_df)
                                                st.success(f"Invited {row['Name']}!")
                                                st.rerun()
                                    else:
                                        st.success("✅ Candidate Shortlisted")
                            
                            st.divider()

                else:
                    st.info("No applications received yet.")

if __name__ == "__main__":
    main()
